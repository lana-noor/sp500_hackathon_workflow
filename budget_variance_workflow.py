# Copyright (c) Microsoft. All rights reserved.

"""
Budget Variance Report Workflow — Six Sequential Agents
========================================================

Pipeline:
  Agent 1  →  MCP Data Agent              (agent_reference: BudgetReportsMCPAgent)
               Retrieves department-submitted variance narrative reports with
               justifications, explanations, and remediation plans.
               Source: MCP server serving department_reports/*.md files

  Agent 2  →  Web Search Agent            (agent_reference: WebSearchBudgetsAgent)
               Searches for macroeconomic context (UAE inflation, sector benchmarks)
               to validate department claims against external economic data.

  Agent 3  →  Code Interpreter Agent      (agent_reference: BudgetVarianceCodeIntAgent)
               Analyzes official CSV files (approved_budgets.csv, historical_actuals.csv)
               and reconciles actual data against department claims from Agent 1.
               Produces structured JSON variance analysis.

  Agent 4  →  Foundry IQ Policy Agent     (agent_reference: BudgetPolicyAgent)
               Uses Azure AI Search (Foundry IQ) to retrieve policy guidance from
               ingested documents (Financial Management Act, procurement guidelines).
               Provides compliance requirements and regulatory context.

  Agent 5  →  Summary Agent               (Responses API - no name)
               Synthesizes outputs from Agents 1-4 into a comprehensive executive
               Markdown report and converts to Word (.docx) document.

  Agent 6  →  Outlook Mail Agent          (agent_reference: BudgetWorkIQMailAgent)
               Sends the final report to lananoor@microsoft.com via Outlook.

Prerequisites:
  - All agent_reference agents deployed in Azure AI Foundry
  - BudgetReportsMCPAgent connected to: https://budget-reports-mcp-server.<id>.eastus.azurecontainerapps.io/mcp
  - BudgetPolicyAgent connected to Azure AI Search index: adga-budget-policies
  - Azure credentials configured (DefaultAzureCredential / AzureCliCredential)
  - Environment variables set (see .envsample)
"""

import asyncio
import json
import re
from datetime import datetime
from pathlib import Path
from typing import List, Optional

import os
from dotenv import load_dotenv
from pydantic import BaseModel

from agent_framework import Agent
from agent_framework.azure import AzureOpenAIResponsesClient
from azure.ai.projects import AIProjectClient
from azure.identity import DefaultAzureCredential

# ---------------------------------------------------------------------------
# python-docx — optional; only needed for the Word executor
# ---------------------------------------------------------------------------
try:
    from docx import Document as WordDocument
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[Warning] python-docx not installed — Word output will be skipped.")
    print("          Install with: pip install python-docx")

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
load_dotenv()

OUTPUT_DIR = Path(os.getenv("OUTPUT_DIR", "./output"))
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

PROMPTS_DIR = Path(__file__).parent / "prompts"

def load_prompt(filename: str) -> str:
    """Load a prompt from the prompts directory."""
    prompt_path = PROMPTS_DIR / filename
    with open(prompt_path, "r", encoding="utf-8") as f:
        return f.read().strip()

MCP_SERVER_URL       = os.getenv("MCP_SERVER_URL", "https://budget-reports-mcp-server.redwave-ed431b4a.eastus.azurecontainerapps.io/mcp")
OUTLOOK_RECIPIENT    = os.getenv("OUTLOOK_RECIPIENT_EMAIL", "lananoor@microsoft.com")

# Agent references (Azure AI Foundry deployed agents)
# Agent 1: BudgetReportsMCPAgent - Department narrative reports via MCP
BUDGET_MCP_AGENT          = os.getenv("BUDGET_MCP_AGENT_NAME", "BudgetReportsMCPAgent")
BUDGET_MCP_VERSION        = os.getenv("BUDGET_MCP_AGENT_VERSION", "3")

# Agent 2: WebSearchBudgetsAgent - Economic validation via web search
WEB_SEARCH_AGENT          = os.getenv("WEB_SEARCH_AGENT_NAME", "WebSearchBudgetsAgent")
WEB_SEARCH_VERSION        = os.getenv("WEB_SEARCH_AGENT_VERSION", "14")

# Agent 3: BudgetVarianceCodeIntAgent - CSV analysis via code interpreter
CODE_INTERPRETER_AGENT    = os.getenv("CODE_INTERPRETER_AGENT_NAME", "BudgetVarianceCodeIntAgent")
CODE_INTERPRETER_VERSION  = os.getenv("CODE_INTERPRETER_AGENT_VERSION", "1")

# Agent 4: BudgetPolicyAgent - Policy guidance via Foundry IQ (AI Search)
POLICY_AGENT              = os.getenv("POLICY_AGENT_NAME", "BudgetPolicyAgent")
POLICY_VERSION            = os.getenv("POLICY_AGENT_VERSION", "1")

# Agent 5: Summary Agent - Uses Responses API (no agent reference)
# Agent 6: BudgetWorkIQMailAgent - Email delivery via Outlook
OUTLOOK_AGENT             = os.getenv("OUTLOOK_AGENT_NAME", "BudgetWorkIQMailAgent")
OUTLOOK_VERSION           = os.getenv("OUTLOOK_AGENT_VERSION", "7")


# ---------------------------------------------------------------------------
# Pydantic structured-output model for Agent 3 (Code Interpreter)
# ---------------------------------------------------------------------------

class DepartmentVariance(BaseModel):
    department_code: str
    department_name: str
    approved_budget_aed: float
    actual_spend_aed: float
    variance_aed: float
    variance_pct: float
    policy_status: str          # ACCEPTABLE | MINOR | SIGNIFICANT | CRITICAL | UNDERSPEND_REVIEW
    required_action: str
    trend_vs_prior_year: Optional[str]   # IMPROVING | WORSENING | STABLE | INSUFFICIENT_DATA


class BudgetVarianceOutput(BaseModel):
    period: str                          # e.g. "Q1 2026"
    organisation: str
    total_approved_aed: float
    total_actual_aed: float
    total_variance_aed: float
    total_variance_pct: float
    overall_policy_status: str
    departments: List[DepartmentVariance]
    departments_requiring_cfo_approval: List[str]
    departments_requiring_board_notification: List[str]
    key_findings: List[str]
    data_quality_notes: List[str]


# ---------------------------------------------------------------------------
# Load Agent Prompts from Files
# ---------------------------------------------------------------------------
MCP_DATA_INSTRUCTIONS = load_prompt("agent1_mcp_data.txt")
# agent2_web_search.txt is the SYSTEM PROMPT configured in the Foundry portal agent.
# We do NOT send it as the user message — the deployed agent already has it.
# We reference it here only for documentation/debugging.
_WEB_SEARCH_SYSTEM_PROMPT_REF = load_prompt("agent2_web_search.txt")
CODE_INTERPRETER_INSTRUCTIONS = load_prompt("agent3_code_interpreter.txt")
POLICY_INSTRUCTIONS = load_prompt("agent4_foundry_iq.txt")
SUMMARY_INSTRUCTIONS = load_prompt("agent5_summary.txt")


# ---------------------------------------------------------------------------
# Executor — Markdown → Word Document (reused from payslip workflow)
# ---------------------------------------------------------------------------

def markdown_to_word_executor(markdown_content: str, output_path: Path) -> Path:
    """Convert a Markdown string to a formatted Word (.docx) document."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc = WordDocument()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    lines = markdown_content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)
        elif stripped.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            data_rows = [
                tl for tl in table_lines
                if not re.fullmatch(r"[\|\-\: ]+", tl.strip())
            ]
            if data_rows:
                parsed: list[list[str]] = []
                for tl in data_rows:
                    cells = [
                        re.sub(r"\*\*(.+?)\*\*", r"\1", c.strip())
                        for c in tl.strip().strip("|").split("|")
                    ]
                    parsed.append(cells)
                num_cols = max(len(r) for r in parsed)
                table = doc.add_table(rows=len(parsed), cols=num_cols)
                table.style = "Table Grid"
                for ri, row_data in enumerate(parsed):
                    for ci in range(num_cols):
                        cell_text = row_data[ci] if ci < len(row_data) else ""
                        cell = table.rows[ri].cells[ci]
                        cell.text = cell_text
                        if ri == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
            continue
        elif re.match(r"^[\-\*] ", stripped):
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped[2:])
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            doc.add_paragraph(text, style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            text = re.sub(r"^\d+\. ", "", stripped)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            doc.add_paragraph(text, style="List Number")
        elif stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
        else:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            if text:
                doc.add_paragraph(text)
        i += 1

    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _extract_text_from_result(result) -> str:
    """Pull the final assistant text from an agent result."""
    try:
        if hasattr(result, "messages") and result.messages:
            for message in reversed(result.messages):
                if hasattr(message, "contents"):
                    for content in message.contents:
                        if getattr(content, "type", None) == "text":
                            if hasattr(content, "text") and content.text:
                                return content.text
    except Exception:
        pass
    return str(result)


def _banner(title: str) -> None:
    bar = "=" * 65
    print(f"\n{bar}")
    print(f"  {title}")
    print(bar)


# ---------------------------------------------------------------------------
# Workflow Executor
# ---------------------------------------------------------------------------

class BudgetVarianceWorkflowExecutor:
    """
    Sequential six-agent workflow executor for Budget Variance Report analysis.

    Step 1  — MCP Data Agent (BudgetReportsMCPAgent)
              Retrieves department-submitted variance narrative reports containing
              justifications, explanations, and remediation plans.
              Source: MCP server serving department_reports/*.md files

    Step 2  — Web Search Agent (WebSearchBudgetsAgent)
              Searches for macroeconomic context (UAE inflation, sector benchmarks)
              to validate department claims against external economic data.

    Step 3  — Code Interpreter Agent (BudgetVarianceCodeIntAgent)
              Analyzes official CSV files (approved_budgets.csv, historical_actuals.csv)
              and reconciles actual data against department claims from Agent 1.
              Performs 3-way reconciliation: Claims vs. Reality vs. External validation.

    Step 4  — Foundry IQ Policy Agent (BudgetPolicyAgent)
              Uses Azure AI Search to retrieve policy guidance from ingested documents
              (Financial Management Act, procurement guidelines, IT spending rules).
              Provides compliance requirements and regulatory context.

    Step 5  — Summary Agent (Responses API)
              Synthesizes outputs from Agents 1-4 into a comprehensive executive
              Markdown report, then converts to Word (.docx).

    Step 6  — Outlook Mail Agent (BudgetWorkIQMailAgent)
              Sends the final report to lananoor@microsoft.com via Outlook.
    """

    def __init__(self) -> None:
        self.responses_client = AzureOpenAIResponsesClient(
            credential=DefaultAzureCredential(),
            project_endpoint=os.getenv("AZURE_AI_PROJECT_ENDPOINT"),
            deployment_name=os.getenv("AZURE_AI_MODEL_DEPLOYMENT_NAME"),
        )
        self.project_client = AIProjectClient(
            endpoint=os.getenv("AZURE_AI_PROJECT_ENDPOINT"),
            credential=DefaultAzureCredential(),
        )
        self.openai_client = self.project_client.get_openai_client()

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    async def execute(self, budget_report_input: str) -> dict:
        """
        Run the full budget variance workflow with 6 agents.

        Args:
            budget_report_input: Budget variance report in Markdown format.

        Returns:
            dict with keys: mcp_data, web_context, analysis, policy_guidance,
                            summary, markdown_report_path, word_doc_path, mail_result.
        """
        _banner("BUDGET VARIANCE WORKFLOW — START (6 AGENTS)")

        # ── Step 1: MCP Data Retrieval (Department Claims) ──────────────
        mcp_data_text = await self._run_mcp_data(budget_report_input)

        # ── Step 2: Web Search Context (Economic Validation) ────────────
        # Pass Agent 1 output so Agent 2 knows WHICH department claims to validate
        web_context_text = await self._run_web_search(budget_report_input, mcp_data_text)

        # ── Step 3: Code Interpreter Analysis (Official CSV Data) ───────
        analysis_text = await self._run_code_analysis(
            budget_report_input, mcp_data_text, web_context_text
        )

        # ── Step 4: Foundry IQ Policy Guidance (AI Search) ──────────────
        policy_text = await self._run_policy_agent(
            mcp_data_text, web_context_text, analysis_text
        )

        # ── Step 5: Summary → Markdown + Word (Responses API) ───────────
        summary_markdown = await self._run_summary(
            budget_report_input, mcp_data_text, web_context_text, analysis_text, policy_text
        )

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        md_path = OUTPUT_DIR / f"budget_variance_report_{timestamp}.md"
        md_path.write_text(summary_markdown, encoding="utf-8")
        print(f"\n[Executor] Markdown saved → {md_path}")

        word_path: Optional[Path] = None
        word_doc_path = OUTPUT_DIR / f"budget_variance_report_{timestamp}.docx"
        try:
            markdown_to_word_executor(summary_markdown, word_doc_path)
            word_path = word_doc_path
            print(f"[Executor] Word document saved → {word_path}")
        except RuntimeError as exc:
            print(f"[Executor] Word conversion skipped: {exc}")

        # ── Step 6: Outlook Mail (Email Delivery) ───────────────────────
        mail_result = await self._run_outlook_mail(summary_markdown, word_path)

        _banner("WORKFLOW COMPLETE")

        return {
            "mcp_data": mcp_data_text,
            "web_context": web_context_text,
            "analysis": analysis_text,
            "policy_guidance": policy_text,
            "summary": summary_markdown,
            "markdown_report_path": str(md_path),
            "word_doc_path": str(word_path) if word_path else None,
            "mail_result": mail_result,
        }

    # ------------------------------------------------------------------
    # Agent 1 — MCP Data Agent (BudgetReportsMCPAgent)
    # ------------------------------------------------------------------

    async def _run_mcp_data(self, budget_report_input: str) -> str:
        print(f"\n[Agent 1] MCP Data Agent — calling {BUDGET_MCP_AGENT} v{BUDGET_MCP_VERSION}...")
        print("[Agent 1] Retrieving department-submitted narrative reports...")

        # Use the prompt instructions from the agent configuration
        response = self.openai_client.responses.create(
            input=[
                {
                    "role": "user",
                    "content": (
                        "Retrieve all department variance narrative reports for Q1 2026. "
                        "Use the MCP tools to gather:\n"
                        "- IT department variance report\n"
                        "- HR department variance report\n"
                        "- Infrastructure department variance report\n\n"
                        "These reports contain department justifications and claims. "
                        "Return the full narrative content.\n\n"
                        f"Context from submitted report:\n{budget_report_input[:500]}"
                    ),
                }
            ],
            extra_body={
                "agent_reference": {
                    "name": BUDGET_MCP_AGENT,
                    "version": str(BUDGET_MCP_VERSION),
                    "type": "agent_reference",
                }
            },
        )
        text = response.output_text
        print(f"[Agent 1] Done. Retrieved department narratives.\n{text[:200]}...\n")
        return text

    # ------------------------------------------------------------------
    # Agent 2 — Web Search Agent (WebSearchBudgetsAgent)
    # ------------------------------------------------------------------

    async def _run_web_search(self, budget_report_input: str, mcp_data_text: str) -> str:
        """
        Agent 2: Web Search — validates department claims from Agent 1 against
        external economic data (UAE inflation, cloud costs, energy, cybersecurity, etc.).

        IMPORTANT: The agent deployed in Foundry already has its system prompt configured.
        We send only the USER MESSAGE here — a focused task with Agent 1's claims as context.
        """
        print(f"\n[Agent 2] Web Search Agent — calling {WEB_SEARCH_AGENT} v{WEB_SEARCH_VERSION}...")
        print("[Agent 2] Validating department claims against Q1 2026 economic data...")

        # Build a focused user-turn message so the Foundry agent knows exactly what to validate.
        # The agent's system prompt (agent2_web_search.txt) already defines search strategy,
        # output schema, and validation language — we must NOT repeat it here.
        user_message = (
            "Validate the following department variance claims against Q1 2026 "
            "(January – March 2026) external economic data.\n\n"
            "=== DEPARTMENT NARRATIVE CLAIMS (Agent 1 / MCP output) ===\n"
            f"{mcp_data_text}\n\n"
            "For each claim above, search for current data and return "
            "CONFIRMED / PARTIALLY CONFIRMED / UNCLEAR / CONTRADICTS.\n"
            "Focus areas: UAE inflation, cloud computing cost trends, energy & utilities pricing, "
            "cybersecurity incident landscape, public-sector recruitment market, and "
            "government regulatory mandates (Zero Trust, procurement rules).\n\n"
            "Return ONLY valid JSON matching the schema in your instructions."
        )

        response = self.openai_client.responses.create(
            input=[{"role": "user", "content": user_message}],
            extra_body={
                "agent_reference": {
                    "name": WEB_SEARCH_AGENT,
                    "version": str(WEB_SEARCH_VERSION),
                    "type": "agent_reference",
                }
            },
        )
        text = response.output_text
        print(f"[Agent 2] Done. Economic validation complete.\n{text[:200]}...\n")
        return text

    # ------------------------------------------------------------------
    # Agent 3 — Code Interpreter Agent (BudgetVarianceCodeIntAgent)
    # ------------------------------------------------------------------

    async def _run_code_analysis(
        self,
        budget_report_input: str,
        mcp_data_text: str,
        web_context_text: str,
    ) -> str:
        print(f"\n[Agent 3] Code Interpreter Agent — calling {CODE_INTERPRETER_AGENT} v{CODE_INTERPRETER_VERSION}...")
        print("[Agent 3] Analyzing official CSV files and reconciling against claims...")
        query = (
            "Use code interpreter to perform 3-way reconciliation analysis.\n\n"
            "You have access to the following CSV files as attachments:\n"
            "- approved_budgets.csv (official approved budgets)\n"
            "- historical_actuals.csv (official spending data)\n\n"
            "=== DEPARTMENT CLAIMS (from MCP) ===\n"
            f"{mcp_data_text}\n\n"
            "=== ECONOMIC VALIDATION (from Web Search) ===\n"
            f"{web_context_text}\n\n"
            "=== ORIGINAL REPORT CONTEXT ===\n"
            f"{budget_report_input[:500]}\n\n"
            "TASK:\n"
            "1. Load and analyze the CSV files\n"
            "2. Calculate actual variances from official data\n"
            "3. Compare department claims vs. actual data (reconciliation)\n"
            "4. Identify discrepancies and assess credibility\n"
            "5. Apply policy thresholds and flag issues\n"
            "6. Return structured JSON matching BudgetVarianceOutput schema."
        )
        response = self.openai_client.responses.create(
            input=[{"role": "user", "content": query}],
            extra_body={
                "agent_reference": {
                    "name": CODE_INTERPRETER_AGENT,
                    "version": str(CODE_INTERPRETER_VERSION),
                    "type": "agent_reference",
                }
            },
        )
        text = response.output_text
        print(f"[Agent 3] Done. Completed reconciliation analysis.\n{text[:200]}...\n")
        return text

    # ------------------------------------------------------------------
    # Agent 4 — Foundry IQ Policy Agent (BudgetPolicyAgent)
    # ------------------------------------------------------------------

    async def _run_policy_agent(
        self,
        mcp_data_text: str,
        web_context_text: str,
        analysis_text: str,
    ) -> str:
        print(f"\n[Agent 4] Foundry IQ Policy Agent — calling {POLICY_AGENT} v{POLICY_VERSION}...")
        print("[Agent 4] Querying AI Search for policy guidance...")
        query = (
            "Use Azure AI Search to retrieve relevant policy guidance for this budget variance situation.\n\n"
            "=== DEPARTMENT CLAIMS ===\n"
            f"{mcp_data_text[:800]}\n\n"
            "=== ECONOMIC CONTEXT ===\n"
            f"{web_context_text[:800]}\n\n"
            "=== VARIANCE ANALYSIS FINDINGS ===\n"
            f"{analysis_text[:800]}\n\n"
            "TASK:\n"
            "Search the policy documents and provide:\n"
            "1. Relevant sections from the Financial Management Act\n"
            "2. IT Technology Spending Guidelines (if IT overspend detected)\n"
            "3. Procurement guidelines (if applicable)\n"
            "4. Variance thresholds and escalation procedures\n"
            "5. Compliance requirements and timelines\n"
            "6. Any regulatory risks or penalties\n\n"
            "Return structured policy guidance with document references."
        )
        response = self.openai_client.responses.create(
            input=[{"role": "user", "content": query}],
            extra_body={
                "agent_reference": {
                    "name": POLICY_AGENT,
                    "version": str(POLICY_VERSION),
                    "type": "agent_reference",
                }
            },
        )
        text = response.output_text
        print(f"[Agent 4] Done. Retrieved policy guidance.\n{text[:200]}...\n")
        return text

    # ------------------------------------------------------------------
    # Agent 5 — Summary Agent (Responses API)
    # ------------------------------------------------------------------

    async def _run_summary(
        self,
        budget_report_input: str,
        mcp_data_text: str,
        web_context_text: str,
        analysis_text: str,
        policy_text: str,
    ) -> str:
        print("\n[Agent 5] Summary Agent (Responses API) — running...")
        print("[Agent 5] Synthesizing comprehensive executive report...")
        agent = Agent(
            client=self.responses_client,
            instructions=SUMMARY_INSTRUCTIONS,
        )
        query = (
            "Generate the full Markdown executive budget variance report using ALL inputs below.\n\n"
            "=== ORIGINAL SUBMITTED REPORT ===\n"
            f"{budget_report_input}\n\n"
            "=== DEPARTMENT NARRATIVE CLAIMS (Agent 1: MCP) ===\n"
            f"{mcp_data_text}\n\n"
            "=== ECONOMIC VALIDATION CONTEXT (Agent 2: Web Search) ===\n"
            f"{web_context_text}\n\n"
            "=== OFFICIAL DATA ANALYSIS & RECONCILIATION (Agent 3: Code Interpreter) ===\n"
            f"{analysis_text}\n\n"
            "=== POLICY & COMPLIANCE GUIDANCE (Agent 4: Foundry IQ) ===\n"
            f"{policy_text}\n\n"
            "Synthesize a comprehensive executive report that includes:\n"
            "1. Department claims and justifications\n"
            "2. Reconciliation findings (claims vs. reality)\n"
            "3. Economic context validation\n"
            "4. Policy compliance assessment\n"
            "5. Executive summary and recommendations"
        )
        result = await agent.run(query)
        text = _extract_text_from_result(result)
        print("[Agent 5] Done. Executive report generated.\n")
        return text

    # ------------------------------------------------------------------
    # Agent 6 — Outlook Mail Agent (BudgetWorkIQMailAgent)
    # ------------------------------------------------------------------

    async def _run_outlook_mail(
        self, summary_markdown: str, word_doc_path: Optional[Path]
    ) -> str:
        print(f"\n[Agent 6] Outlook Mail Agent — calling {OUTLOOK_AGENT} v{OUTLOOK_VERSION}...")
        print(f"[Agent 6] Sending email to {OUTLOOK_RECIPIENT}...")

        word_note = (
            f"A Word document version of this report has been saved to: {word_doc_path}"
            if word_doc_path
            else "Note: Word document generation was skipped (python-docx not available)."
        )

        email_prompt = f"""
Please compose and send a professional email with the following details:

To: {OUTLOOK_RECIPIENT}
Subject: Q1 2026 Budget Variance Analysis Report — Emirates Digital Authority (EDA)

IMPORTANT: Send this email to {OUTLOOK_RECIPIENT} — this is the confirmed recipient.

Email body should:
1. Open with a brief professional introduction (2-3 sentences) explaining this is an
   AI-assisted budget variance analysis for the Emirates Digital Authority (EDA) for Q1 2026.
2. Include a concise summary of the headline findings:
   - Total overspend amount (AED) and percentage
   - Number of CRITICAL departments (requiring Board notification)
   - Number of SIGNIFICANT departments (requiring CFO approval)
3. List the required actions with their deadlines.
4. State that the full analysis report is included below.
5. Close professionally.

Set the email importance to HIGH if any CRITICAL or SIGNIFICANT variances are present.

Then append the full report content below the email closing:

{summary_markdown}

---
{word_note}
This email was generated automatically by the Budget Variance Workflow (Azure AI Foundry).
""".strip()

        response = self.openai_client.responses.create(
            input=[{"role": "user", "content": email_prompt}],
            extra_body={
                "agent_reference": {
                    "name": OUTLOOK_AGENT,
                    "version": str(OUTLOOK_VERSION),
                    "type": "agent_reference",
                }
            },
        )
        text = response.output_text
        print(f"[Agent 6] Done. Email sent successfully.\n{text[:150]}...\n")
        return text


# ---------------------------------------------------------------------------
# Synthetic sample report — used when the script is run directly
# ---------------------------------------------------------------------------

def _load_sample_report() -> str:
    """Load the sample budget report from the data/ subfolder."""
    # Primary location: data/sample_budget_report.md
    sample_path = Path(__file__).parent / "data" / "sample_budget_report.md"
    if sample_path.exists():
        return sample_path.read_text(encoding="utf-8")
    # Fallback: same directory as the script (legacy location)
    fallback_path = Path(__file__).parent / "sample_budget_report.md"
    if fallback_path.exists():
        return fallback_path.read_text(encoding="utf-8")
    return "# Budget Variance Report\n\n(sample_budget_report.md not found — expected at data/sample_budget_report.md)"


SAMPLE_BUDGET_REPORT_MARKDOWN = _load_sample_report()


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

async def main() -> None:
    """Run the budget variance workflow on the sample report."""
    executor = BudgetVarianceWorkflowExecutor()
    result = await executor.execute(SAMPLE_BUDGET_REPORT_MARKDOWN)

    _banner("OUTPUT PATHS")
    print(f"Markdown  → {result['markdown_report_path']}")
    print(f"Word Doc  → {result['word_doc_path'] or 'Not generated'}")
    print(f"Mail Sent → {result['mail_result'][:100]}...")

    print("\n--- SUMMARY PREVIEW (first 1000 chars) ---")
    print(result["summary"][:1000])


if __name__ == "__main__":
    asyncio.run(main())
