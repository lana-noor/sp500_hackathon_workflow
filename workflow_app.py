# Copyright (c) Microsoft. All rights reserved.

import asyncio
import json
import re
from datetime import datetime
from pathlib import Path
from typing import List, Optional

from pydantic import BaseModel

from agent_framework import Agent
from agent_framework.azure import AzureOpenAIResponsesClient, AzureAIAgentClient
import os
from dotenv import load_dotenv
from azure.identity.aio import DefaultAzureCredential, AzureCliCredential

"""
Payslip Verification Workflow — Three Sequential Agents + Word Executor

Pipeline:
  Agent 1  →  Document Verification Agent   (validates structure & fraud signals)
  Agent 2  →  Salary Analysis Agent          (code interpreter + CSV, structured output)
  Agent 3  →  Document Summary Agent         (produces Markdown report)
  Executor →  markdown_to_word_executor()    (converts Markdown → .docx)
"""

# ---------------------------------------------------------------------------
# python-docx — optional; only needed for the Word executor
# ---------------------------------------------------------------------------
try:
    from docx import Document as WordDocument
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[Warning] python-docx not installed — Word output will be skipped.")
    print("          Install with: pip install python-docx")

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
load_dotenv()  # Load environment variables from .env file
PAYSLIPS_CSV = os.getenv("PAYSLIPS_CSV_PATH")
OUTPUT_DIR = Path(os.getenv("OUTPUT_DIR"))

# ---------------------------------------------------------------------------
# Pydantic structured-output model for Agent 2
# ---------------------------------------------------------------------------

class IncomeEntry(BaseModel):
    date: Optional[str]
    income: Optional[float]
    income_type: Optional[str]


class VariationPair(BaseModel):
    absolute: Optional[float]
    percent: Optional[float]


class VariationAnalysis(BaseModel):
    from_1_to_2: VariationPair
    from_2_to_3: VariationPair
    overall_1_to_3: VariationPair


class SalaryAnalysisOutput(BaseModel):
    employee_name: Optional[str]
    income_by_date: List[IncomeEntry]
    variation: VariationAnalysis
    valid: bool
    validation_reasoning: List[str]


# ---------------------------------------------------------------------------
# Agent 1 — Document Verification Agent
# ---------------------------------------------------------------------------
VERIFICATION_INSTRUCTIONS = """
You are a Document Verification Agent specialized in salary slip fraud detection
and authenticity checks for UAE-based payslips.

YOUR TASK
Analyze the provided salary slip document (Markdown or JSON format) and determine:

1. REQUIRED FIELDS CHECK — verify that all of these are present:
   - Company Name, Employee Name, Employee ID, Designation
   - Pay Period / Pay Date
   - Basic Salary, Housing Allowance
   - Gross Earnings, Net Salary

2. MATHEMATICAL CONSISTENCY:
   - Gross Earnings must equal the sum of all individual earnings components.
   - Net Salary must equal Gross Earnings minus Total Deductions.
   - Flag any discrepancy (tolerance: ±1 AED for rounding).

3. SPELLING & GRAMMAR:
   - Check all visible text fields for spelling errors.
   - Flag incorrect designations, misspelled field labels, or garbled text.

4. FORMAT VALIDATION (UAE payslip conventions):
   - Currency should be AED.
   - Dates should be in a recognisable format (YYYY-MM-DD or Month YYYY).
   - Allowances and deductions must be itemised, not lump-sum only.

5. FRAUD INDICATORS:
   - Math errors (gross ≠ sum of earnings; net ≠ gross − deductions).
   - Implausibly round numbers for every field simultaneously.
   - Missing employer details.
   - Employee name inconsistency across fields.
   - Unusually high overtime (>50% of basic) with no stated reason.
   - Zero salary with no explanation.

OUTPUT FORMAT
Return ONLY a valid JSON object — no prose, no markdown fences:
{
  "document_valid": true or false,
  "employee_name": "full name extracted from slip, or null",
  "employee_id": "ID extracted from slip, or null",
  "verification_checks": {
    "fields_complete": true or false,
    "math_consistent": true or false,
    "format_valid": true or false,
    "no_fraud_flags": true or false
  },
  "issues_found": ["list each issue as a string; empty list if none"],
  "verification_summary": "one concise paragraph summarising the result"
}
""".strip()

# ---------------------------------------------------------------------------
# Agent 2 — Salary Analysis Agent (code interpreter + structured output)
# ---------------------------------------------------------------------------
ANALYSIS_INSTRUCTIONS = """
You are a Salary Slip Analysis Agent with Code Interpreter enabled.

INPUT
- An employee name to look up.
- A CSV file (payslips_batch_manifest.csv) is attached to the code interpreter.
- Optionally, the raw submitted salary slip for cross-reference.

TASK
1) Use code interpreter to read the CSV and filter all rows where
   employee_name exactly matches the provided name.
2) For each matching row extract:
   - employee_name  (verify consistency)
   - salary_date   (use pay_date column; normalize to YYYY-MM-DD)
   - total_income  (prefer net_salary_aed; else gross_earnings_aed;
                   else compute: gross_earnings_aed - deductions_aed)
3) Normalize all numbers (strip commas/currency symbols; treat blanks as 0).
4) Sort by salary_date ascending (oldest → newest).
5) Compute variations:
   - Absolute and % change: slip1→slip2, slip2→slip3 (if present).
   - Overall absolute and % change: slip1→slip3 (if ≥3 slips).
6) Apply VARIATION VALIDATION THRESHOLDS (MANDATORY):
   - |% change| < 20%      → APPROVE  (valid=true) unless other issues exist.
   - 20% ≤ |% change| < 40% → NEEDS CLARIFICATION (valid=false).
     * Set valid=true ONLY if the slip text explicitly states a reason
       (bonus, arrears, unpaid leave, large deduction) AND the math aligns.
   - |% change| ≥ 40%      → DECLINE (valid=false).
     * Override to valid=true ONLY if explicit explanation + amounts reconcile.
7) Also compare submitted slip values against CSV values; flag any mismatch.

VALIDATION REASONING must reference:
- % change computed and the threshold bucket applied.
- Whether an explicit explanation was found in the slip (quote the field/line).
- Why the result is approved / flagged / declined.
- Any currency mismatch, missing date, or gross vs net inconsistency.

OUTPUT RULES
- Return ONLY valid JSON — no prose, no markdown fences.
- Adhere strictly to the schema below.
- If only 1 slip found: all variation fields must be null.

REQUIRED JSON OUTPUT SCHEMA
{
  "employee_name": "string or null",
  "income_by_date": [
    { "date": "YYYY-MM-DD or null", "income": number or null, "income_type": "net|gross|computed|null" }
  ],
  "variation": {
    "from_1_to_2": { "absolute": number or null, "percent": number or null },
    "from_2_to_3": { "absolute": number or null, "percent": number or null },
    "overall_1_to_3": { "absolute": number or null, "percent": number or null }
  },
  "valid": true or false,
  "validation_reasoning": ["string", "string"]
}
""".strip()

# ---------------------------------------------------------------------------
# Agent 3 — Document Summary Agent
# ---------------------------------------------------------------------------
SUMMARY_INSTRUCTIONS = """
You are a Document Summary Agent. Synthesise the outputs of a document
verification check and a salary analysis into a professional Markdown report.

You will receive:
  1. Document Verification Result  (JSON from Agent 1)
  2. Salary Analysis Result        (JSON from Agent 2)

Generate a Markdown report that follows this EXACT structure:

---

# Salary Slip Verification Report

**Generated:** {today's date}
**Employee:** {employee_name}

---

## Executive Summary

{1–2 sentence summary of overall result and approval status}

---

## 1. Document Verification

### Status: ✅ APPROVED / ❌ REJECTED / ⚠️ NEEDS REVIEW

| Check | Result |
|---|---|
| Fields Complete | ✅ Yes / ❌ No |
| Math Consistent | ✅ Yes / ❌ No |
| Format Valid | ✅ Yes / ❌ No |
| No Fraud Flags | ✅ Yes / ❌ No |

**Issues Found:**
{bulleted list; or "No issues found." if empty}

---

## 2. Salary Analysis

### Salary Data

| Pay Date | Net Income (AED) | Income Type |
|---|---|---|
{one row per slip, sorted oldest to newest}

### Variation Analysis

| Period | Absolute Change (AED) | % Change | Decision |
|---|---|---|---|
{one row per pair: Slip 1→2, Slip 2→3, Overall}

### Validation Status: ✅ VALID / ❌ INVALID

**Reasoning:**
{numbered list of validation_reasoning points}

---

## 3. Final Decision

| Field | Value |
|---|---|
| Employee | {name} |
| Overall Status | APPROVED / REJECTED / NEEDS REVIEW |
| Document Valid | ✅ Yes / ❌ No |
| Salary Valid | ✅ Yes / ❌ No |
| Recommendation | {one clear sentence} |

---

## 4. Notes & Observations

{any additional context, warnings, or caveats — or "None." if nothing to add}

---

*Report generated by Payslip Verification Workflow*

---

DECISION RULES:
- APPROVED     → both document_valid AND salary valid are true.
- REJECTED     → either is definitively false with no path to resolution.
- NEEDS REVIEW → clarification required (e.g., 20–40% salary variation).
""".strip()

# ---------------------------------------------------------------------------
# Executor — Markdown → Word Document
# ---------------------------------------------------------------------------

def markdown_to_word_executor(markdown_content: str, output_path: Path) -> Path:
    """
    Convert a Markdown string to a formatted Word (.docx) document.
    Handles: headings (H1–H4), tables, bullet lists, bold text, and HR lines.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    doc = WordDocument()

    # Set default body style
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    lines = markdown_content.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Empty lines → paragraph break
        if not stripped:
            i += 1
            continue

        # H1 – H4 headings
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)

        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)

        # Markdown table block — collect all consecutive pipe-lines
        elif stripped.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            # Filter out separator rows (|---|---|)
            data_rows = [
                tl for tl in table_lines
                if not re.fullmatch(r"[\|\-\: ]+", tl.strip())
            ]

            if data_rows:
                parsed: list[list[str]] = []
                for tl in data_rows:
                    cells = [
                        re.sub(r"\*\*(.+?)\*\*", r"\1", c.strip())
                        for c in tl.strip().strip("|").split("|")
                    ]
                    parsed.append(cells)

                num_cols = max(len(r) for r in parsed)
                table = doc.add_table(rows=len(parsed), cols=num_cols)
                table.style = "Table Grid"
                for ri, row_data in enumerate(parsed):
                    for ci in range(num_cols):
                        cell_text = row_data[ci] if ci < len(row_data) else ""
                        cell = table.rows[ri].cells[ci]
                        cell.text = cell_text
                        if ri == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
            continue  # i was already advanced inside the while loop

        # Unordered bullet points
        elif re.match(r"^[\-\*] ", stripped):
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped[2:])
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            doc.add_paragraph(text, style="List Bullet")

        # Numbered list
        elif re.match(r"^\d+\. ", stripped):
            text = re.sub(r"^\d+\. ", "", stripped)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            doc.add_paragraph(text, style="List Number")

        # Bold-only line (standalone **text**)
        elif stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True

        # Regular paragraph
        else:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            if text:
                doc.add_paragraph(text)

        i += 1

    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _extract_text_from_result(result) -> str:
    """
    Pull the final assistant text from an agent result.
    Tries message contents first; falls back to str(result).
    """
    try:
        if hasattr(result, "messages") and result.messages:
            for message in reversed(result.messages):
                if hasattr(message, "contents"):
                    for content in message.contents:
                        if getattr(content, "type", None) == "text":
                            if hasattr(content, "text") and content.text:
                                return content.text
    except Exception:
        pass
    return str(result)


def _extract_employee_name(verification_text: str) -> Optional[str]:
    """Parse employee_name from the verification agent's JSON output."""
    # Use greedy match to capture the outermost JSON object (not a nested one)
    try:
        start = verification_text.find("{")
        end = verification_text.rfind("}")
        if start != -1 and end != -1 and end > start:
            data = json.loads(verification_text[start:end + 1])
            name = data.get("employee_name")
            if name and name.strip().lower() not in ("null", "none", ""):
                return name.strip()
    except (json.JSONDecodeError, AttributeError):
        pass
    return None


# ---------------------------------------------------------------------------
# Workflow Executor
# ---------------------------------------------------------------------------

class PayslipWorkflowExecutor:
    """
    Sequential three-agent workflow executor.

    Step 1  — Document Verification Agent
              Validates the submitted salary slip for completeness,
              mathematical consistency, formatting, and fraud signals.
              Extracts the employee name.

    Step 2  — Salary Analysis Agent
              Uses code interpreter with the payslips CSV to retrieve the
              employee's historical salary data, compute net-pay variations,
              and validate them against defined thresholds.
              Returns structured JSON (SalaryAnalysisOutput schema).

    Step 3  — Document Summary Agent
              Combines both results into a professionally formatted
              Markdown verification report.

    Executor — markdown_to_word_executor()
              Converts the Markdown report to a Word (.docx) document.

    References:
        https://learn.microsoft.com/en-us/agent-framework/workflows/executors
    """

    def __init__(self) -> None:
        self.client = AzureOpenAIResponsesClient(
            credential=DefaultAzureCredential(),
            project_endpoint=os.getenv("AZURE_AI_PROJECT_ENDPOINT"),
            deployment_name=os.getenv("AZURE_AI_MODEL_DEPLOYMENT_NAME"),
        )

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    async def execute(self, salary_slip_input: str) -> dict:
        """
        Run the full payslip verification workflow.

        Args:
            salary_slip_input: Salary slip content in Markdown or JSON format.

        Returns:
            dict with keys:
              - verification       (str)
              - analysis           (str)
              - summary            (str)
              - markdown_report_path (str)
              - word_doc_path      (str | None)
        """
        _banner("PAYSLIP VERIFICATION WORKFLOW — START")

        # ── Step 1: Document Verification ──────────────────────────────
        verification_text = await self._run_verification(salary_slip_input)
        employee_name = _extract_employee_name(verification_text)

        if not employee_name:
            print("[Agent 1] ⚠ Could not extract employee name; defaulting to 'Unknown'.")
            employee_name = "Unknown Employee"
        else:
            print(f"[Agent 1] Employee identified: {employee_name}")

        # ── Step 2: Salary Analysis ─────────────────────────────────────
        analysis_text = await self._run_analysis(salary_slip_input, employee_name)

        # ── Step 3: Summary ─────────────────────────────────────────────
        summary_markdown = await self._run_summary(verification_text, analysis_text)

        # ── Executor: Save Markdown ─────────────────────────────────────
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = re.sub(r"[^\w]", "_", employee_name)

        md_path = OUTPUT_DIR / f"payslip_report_{safe_name}_{timestamp}.md"
        md_path.write_text(summary_markdown, encoding="utf-8")
        print(f"\n[Executor] Markdown saved → {md_path}")

        # ── Executor: Markdown → Word ────────────────────────────────────
        word_path: Optional[Path] = None
        word_doc_str = OUTPUT_DIR / f"payslip_report_{safe_name}_{timestamp}.docx"
        try:
            markdown_to_word_executor(summary_markdown, word_doc_str)
            word_path = word_doc_str
            print(f"[Executor] Word document saved → {word_path}")
        except RuntimeError as exc:
            print(f"[Executor] Word conversion skipped: {exc}")

        _banner("WORKFLOW COMPLETE")

        return {
            "verification": verification_text,
            "analysis": analysis_text,
            "summary": summary_markdown,
            "markdown_report_path": str(md_path),
            "word_doc_path": str(word_path) if word_path else None,
        }

    # ------------------------------------------------------------------
    # Private agent runners
    # ------------------------------------------------------------------

    async def _run_verification(self, salary_slip_input: str) -> str:
        print("\n[Agent 1] Document Verification Agent — running...")
        agent = Agent(
            client=self.client,
            instructions=VERIFICATION_INSTRUCTIONS,
        )
        result = await agent.run(
            f"Please verify the following salary slip document and return the JSON result:\n\n"
            f"{salary_slip_input}"
        )
        text = _extract_text_from_result(result)
        print(f"[Agent 1] Done.\n{text[:200]}...\n")
        return text

    async def _run_analysis(self, salary_slip_input: str, employee_name: str) -> str:
        print(f"[Agent 2] Salary Analysis Agent (Azure AI Agent Service) — running for: {employee_name}...")
        # Read CSV content to embed in the prompt; code interpreter will parse it with Python
        csv_content = ""
        if PAYSLIPS_CSV:
            try:
                csv_content = Path(PAYSLIPS_CSV).read_text(encoding="utf-8")
            except OSError as exc:
                csv_content = f"[Could not read CSV: {exc}]"
        query = (
            f"Use code interpreter to complete this analysis.\n\n"
            f"TARGET EMPLOYEE: {employee_name}\n\n"
            f"STEP 1 — Run Python code to parse the CSV string below with io.StringIO + csv.DictReader:\n"
            f"```csv\n{csv_content}\n```\n\n"
            f"STEP 2 — Filter ALL rows where employee_name == '{employee_name}'. "
            f"There will be multiple rows (multiple months). Do NOT stop at the first match.\n\n"
            f"STEP 3 — For each matching row extract: pay_date, net_salary_aed.\n\n"
            f"STEP 4 — Sort by pay_date ascending and compute variation analysis per instructions.\n\n"
            f"STEP 5 — Return ONLY valid JSON matching the required output schema — no prose, no markdown fences.\n\n"
            f"=== SUBMITTED SALARY SLIP (for cross-reference only) ===\n{salary_slip_input}"
        )
        async with AzureCliCredential() as credential:
            async with AzureAIAgentClient(
                credential=credential,
                project_endpoint=os.getenv("AZURE_AI_PROJECT_ENDPOINT"),
            ).as_agent(
                name="SalaryAnalysisAgent",
                instructions=ANALYSIS_INSTRUCTIONS,
                tools=[AzureAIAgentClient.get_code_interpreter_tool()],
            ) as agent:
                result = await agent.run(query)
        text = _extract_text_from_result(result)
        print(f"[Agent 2] Done.\n{text[:200]}...\n")
        return text

    async def _run_summary(
        self, verification_text: str, analysis_text: str
    ) -> str:
        print("[Agent 3] Document Summary Agent — running...")
        agent = Agent(
            client=self.client,
            instructions=SUMMARY_INSTRUCTIONS,
        )
        query = (
            f"Generate the full Markdown verification report using these inputs.\n\n"
            f"=== DOCUMENT VERIFICATION RESULT ===\n{verification_text}\n\n"
            f"=== SALARY ANALYSIS RESULT ===\n{analysis_text}"
        )
        result = await agent.run(query)
        text = _extract_text_from_result(result)
        print(f"[Agent 3] Done.\n")
        return text


# ---------------------------------------------------------------------------
# Utility
# ---------------------------------------------------------------------------

def _banner(title: str) -> None:
    bar = "=" * 60
    print(f"\n{bar}")
    print(f"  {title}")
    print(bar)


# ---------------------------------------------------------------------------
# Sample salary slip — used when the script is run directly
# ---------------------------------------------------------------------------
SAMPLE_SALARY_SLIP_MARKDOWN = """
# Salary Slip — March 2026

**Company:** Digital Transformation Corp
**Employee Name:** Aarav Mehta
**Employee ID:** EMP1001
**Designation:** Operations Analyst
**Pay Period:** March 2026
**Pay Date:** 2026-03-31

## Earnings

| Component            | Amount (AED) |
|----------------------|-------------|
| Basic Salary         | 6,000.00    |
| Housing Allowance    | 1,157.86    |
| Transport Allowance  | 550.00      |
| Food Allowance       | 150.00      |
| Mobile Allowance     | 100.00      |
| Special Allowance    | 150.00      |
| Overtime             | 250.00      |
| **Gross Earnings**   | **8,357.86**|

## Deductions

| Component            | Amount (AED) |
|----------------------|-------------|
| Total Deductions     | 0.00        |

## Net Salary: AED 8,357.86
""".strip()


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

async def main() -> None:
    """Run the payslip verification workflow on the sample salary slip."""
    executor = PayslipWorkflowExecutor()
    result = await executor.execute(SAMPLE_SALARY_SLIP_MARKDOWN)

    _banner("OUTPUT PATHS")
    print(f"Markdown → {result['markdown_report_path']}")
    print(f"Word Doc → {result['word_doc_path'] or 'Not generated (python-docx missing)'}")

    print("\n--- SUMMARY PREVIEW (first 800 chars) ---")
    print(result["summary"][:800])


if __name__ == "__main__":
    asyncio.run(main())
